from datetime import datetime as dt
import locale

from docx import Document as word_doc
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


def create_doc(news_data, filename, date_string):
    doc = initialize_doc(news_data, date_string)
    doc.save(filename)


def initialize_doc(news_data, date_string):
    """Create a new doc. and fill it"""
    doc = create_configure_doc()  # create doc and configure it (margins)
    add_heading(doc, date_string)  # add heading to doc
    for card_dict in news_data:  # fill doc with news data
        add_text_par(doc, card_dict)
    return doc


def create_configure_doc():
    """Create doc and configure margins"""
    doc = word_doc()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(1)
    return doc


def add_heading(doc, date_string):
    heading_par = doc.add_paragraph()
    configure_par(heading_par, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    configure_par_format(heading_par.paragraph_format)
    heading_run = heading_par.add_run(
        f"Политическое информирование за {date_string} г."
    )
    configure_run(heading_run, bold=True)
    empty_par = doc.add_paragraph()
    configure_par_format(empty_par)


def add_text_par(doc, card_dict):
    news_card_par = doc.add_paragraph()
    configure_par(news_card_par)
    configure_par_format(news_card_par.paragraph_format)
    # date text
    date_run = news_card_par.add_run(card_dict["Date_title"])
    configure_run(date_run, bold=True)
    # news text
    news_text_par = doc.add_paragraph()
    configure_par(news_text_par)
    configure_par_format(news_text_par.paragraph_format)
    news_text_run = news_text_par.add_run(card_dict["Title"])
    configure_run(news_text_run)


def configure_par(
    par,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
):
    par.alignment = alignment


def configure_par_format(
    par_format, line_spacing_rule=WD_LINE_SPACING.SINGLE, space_before=0, space_after=0
):
    par_format.line_spacing_rule = line_spacing_rule
    par_format.space_before = Pt(space_before)
    par_format.space_after = Pt(space_after)


def configure_run(
    run,
    font_size=15,
    font="Times New Roman",
    line_spacing=12,
    bold=False,
):
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font
